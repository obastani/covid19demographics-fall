# https://www.mass.gov/info-details/covid-19-cases-quarantine-and-monitoring

from bs4 import BeautifulSoup
import csv
from datetime import datetime
from docx import Document
from io import StringIO
import os
import re
import requests

def run_MA(args):
    # Parameters
    raw_name = '../MA/raw'
    now = str(datetime.now())

    # Step 1: grab the main website to get quarantine counts and the link to the docx
    out = {}
    r = requests.get("https://www.mass.gov/info-details/covid-19-response-reporting").text
    with open("%s/%s.html" % (raw_name, now), "w") as fp:
        fp.write(r)
    soup = BeautifulSoup(r, "html.parser")
    links = soup.find_all("a")
    doclink = [x for x in links if "dashboard" in x.text.lower() and not "inmate" in x.text.lower()]
    if len(doclink) != 1:
        raise Exception("Malformed link")
    docurl = "https://www.mass.gov" + doclink[0]["href"]

    r = requests.get(docurl)
    with open("%s/%s.pdf" % (raw_name, now), "wb") as fp:
        fp.write(r.content)

def run_MA_old(args):
    # Parameters
    raw_name = '../MA/raw'
    data_name = '../MA/data/data.csv'
    now = str(datetime.now())

    # Step 1: grab the main website to get quarantine counts and the link to the docx
    out = {}
    r = requests.get("https://www.mass.gov/info-details/covid-19-cases-quarantine-and-monitoring").text
    with open("%s/%s.html" % (raw_name, now), "w") as fp:
        fp.write(r)
    soup = BeautifulSoup(r, "html.parser")
    links = soup.find_all("a")
    doclink = [x for x in links if "doc" in x.text.lower() and x.has_attr("href") and "cases" in x["href"]]
    if len(doclink) != 1:
        raise Exception("Malformed link")
    docurl = "https://www.mass.gov" + doclink[0]["href"]
    qtab = [x for x in soup.find_all("table") if x.thead and "quarantine" in x.thead.text.lower()]
    if len(qtab) != 1:
        raise Exception("Missing quarantine table")
    redate = re.compile(r"as of (\S+ \d+, \d+)")
    match = redate.search(qtab[0].thead.text)
    if not match:
        raise Exception("Malformed quarantine update date")
    out["QuarantineUpdate"] = datetime.strptime(match.group(1), "%B %d, %Y").strftime("%Y-%m-%d")
    q1 = [x for x in qtab[0].find_all("tr") if "subject to quarantine" in x.th.text]
    q2 = [x for x in qtab[0].find_all("tr") if "completed monitoring" in x.th.text]
    q3 = [x for x in qtab[0].find_all("tr") if "undergoing monitoring" in x.th.text]
    if len(q1) != 1 or len(q2) != 1 or len(q3) != 1:
        raise Exception("Malformed quarantine table")
    out["QuarantineEver"] = int(q1[0].td.text.replace(",",""))
    out["QuarantineComplete"] = int(q2[0].td.text.replace(",",""))
    out["QuarantineCurrent"] = int(q3[0].td.text.replace(",",""))

    # Step 2: Grab the docx and extract age and gender info
    r = requests.get(docurl, stream=True)
    rawpos = "%s/%s.docx" % (raw_name, now)
    with open(rawpos, "wb") as fp:
        for chunk in r.iter_content(1024*1024*2):
            fp.write(chunk)
    doc = Document(rawpos)
    tab = doc.tables[0]  # First table
    header = "top"
    tabrows = {"top": []}  # header -> list of row data
    for row in tab.rows:
        if len(row.cells) != 2:
            raise Exception("Unexpected table width")
        left = "".join([x.text for x in row.cells[0].paragraphs])
        right = "".join([x.text for x in row.cells[1].paragraphs])
        if left == right:
            header = left
            tabrows[header] = []
        else:
            tabrows[header].append((left, right))
    if not "Sex" in tabrows or not "Age Group" in tabrows:
        raise Exception("Missing a table section")
    expectSex = ["Male", "Female", "Unknown"]
    ageMap = {'≤19 years of age': "0_19", '20-29 years of age': "20_29", '30-39 years of age': "30_39", '40-49 years of age': "40_49", '50-59 years of age': "50_59", '60-69 years of age': "60_69", '70-79 years of age': "70_79", '≥ 80 years of age': "80_plus", "Unknown": "Unknown"}
    expectAge = [x for x in ageMap if x != "Unknown"]
    if sorted([x[0] for x in tabrows["Sex"]]) != sorted(expectSex):
        raise Exception("Unexpected sex categories")
    if sorted([x[0] for x in tabrows["Age Group"] if x[0] != "Unknown"]) != sorted(expectAge):
        raise Exception("Unexpected age groups")
    for lab, val in tabrows["Sex"]:
        out["CaseSex" + lab] = int(val)
    for lab, val in tabrows["Age Group"]:
        out["CaseAge" + ageMap[lab]] = int(val)
    if not "CaseAgeUnknown" in out:
        out["CaseAgeUnknown"] = None
    redoctime = re.compile("As of (\S+\s+\d+,\s+\d+)")
    matches = [redoctime.match(x.text.strip()) for x in doc.paragraphs]
    matches = [x for x in matches if x is not None]
    if len(matches) != 1:
        raise Exception("Unexpected update time information for doc")
    out["DocUpdateDate"] = datetime.strptime(matches[0].group(1), "%B %d, %Y").strftime("%Y-%m-%d")

    # Step 3: output to csv
    out["Scrape_Time"] = now
    fields = sorted([x for x in out])
    exists = os.path.exists(data_name)
    with open(data_name, "a") as fp:
        writer = csv.writer(fp)
        if not exists:
            writer.writerow(fields)
        writer.writerow([out[x] for x in fields])

if __name__ == '__main__':
    run_MA({})
